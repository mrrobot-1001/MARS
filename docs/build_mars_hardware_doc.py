from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/MARS_Practical_Hardware_Software_Integration_Plan.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(18, 37, 63)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color="D9E2EC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_cell_text(cell, bold=False, color=INK, size=9.5):
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.1
        for run in p.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_para(doc, text="", style=None, bold=False, color=None, size=None, italic=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size or 11, color=color or INK, bold=bold, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, value in [
        ("w:start", "1"),
        ("w:numFmt", "decimal"),
        ("w:lvlText", "%1."),
        ("w:lvlJc", "left"),
    ]:
        node = OxmlElement(tag)
        node.set(qn("w:val"), value)
        lvl.append(node)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def add_numbered(doc, text, num_id):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = num_pr.get_or_add_ilvl()
    ilvl.val = 0
    num = num_pr.get_or_add_numId()
    num.val = num_id
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="D9E2EC", size="4")
    cell = table.cell(0, 0)
    set_cell_fill(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    run2 = p2.add_run(body)
    set_run_font(run2, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    header_tr_pr.append(tbl_header)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_fill(cell, header_fill)
        cell.text = header
        style_cell_text(cell, bold=True, color=DARK_BLUE, size=9.2)
    for row in rows:
        table_row = table.add_row()
        tr_pr = table_row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        cells = table_row.cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            set_cell_fill(cells[i], WHITE)
            style_cell_text(cells[i], size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "MARS Railway Safety System"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "Practical hardware, integration, and training plan"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("MARS")
    set_run_font(r, size=28, color=BLUE, bold=True)
    r2 = p.add_run(" Railway Safety System")
    set_run_font(r2, size=28, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("Practical Hardware, Software Integration, and Model Training Plan")
    set_run_font(run, size=15, color=MUTED)

    rows = [
        ("Purpose", "Feasible and cost-conscious plan for collecting field data and operating MARS end to end."),
        ("Prepared for", "MARS project implementation and demonstration"),
        ("Current software state", "FastAPI inference service, React dashboard, regional simulator, sensor simulation script, and Hugging Face Space demo"),
        ("Date", date.today().strftime("%B %d, %Y")),
    ]
    add_table(doc, ["Item", "Detail"], rows, [1800, 7560], header_fill=LIGHT_GRAY)
    add_callout(
        doc,
        "Recommended approach",
        "Start with a low-cost pilot on non-critical demonstration track or lab rig, validate sensor quality and model behavior, then harden enclosures, communications, and operating procedures before any field deployment near live railway operations.",
    )


def build_doc():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "MARS is now a region-aware multimodal railway safety system. The software can ingest track sensor telemetry, weather conditions, and video/security signals; apply Indian Railway zone and division context; and return risk severity, recommended action, and explainable feature contributions.",
    )
    add_para(
        doc,
        "The most practical implementation is a phased architecture: inexpensive sensor nodes collect vibration and temperature; a local edge gateway validates and buffers data; optional camera edge AI handles trespasser or obstruction events; and the MARS API performs risk inference and publishes alerts to dashboards or downstream control systems.",
    )

    doc.add_page_break()
    add_heading(doc, "2. What Has Been Implemented in MARS", 1)
    rows = [
        ("Regional profiles", "Added zone profiles for NR, CR, WR, ER, SR, NFR, and NWR with divisions, hazards, and operating factors."),
        ("Backend scoring", "Rule-based track and weather models now adjust for regional speed, vibration, maintenance, and weather sensitivity."),
        ("API contract", "Segment metadata supports optional division; predictions include region, division, and the resolved regional profile."),
        ("Discovery endpoint", "Added GET /regions so clients can list supported railway zones and operating profiles."),
        ("Sensor simulator", "scripts/simulate_sensors.py now generates region-aware track and weather events and calls both inference endpoints."),
        ("Frontend", "React dashboard and evaluator include zone/division controls, regional context, and corrected explanation rendering."),
        ("Hugging Face Space", "Updated to a Gradio multimodal simulator with track, weather, video/security, and region controls."),
        ("Tests", "Added unit tests for region aliases and region-driven rule-based score differences."),
    ]
    add_table(doc, ["Area", "Completed work"], rows, [2100, 7260])

    add_heading(doc, "3. End-to-End Working Architecture", 1)
    add_para(doc, "The production architecture should stay modular so low-cost pilots and hardened field deployments use the same software contract.")
    architecture_num = create_decimal_numbering(doc)
    add_numbered(doc, "Track-side sensor node samples vibration, temperature, optional strain/displacement, GPS/time, and local health data.", architecture_num)
    add_numbered(doc, "Node publishes compact JSON or MQTT messages to an edge gateway over wired Ethernet, RS-485, Wi-Fi, LoRaWAN, or cellular depending on the site.", architecture_num)
    add_numbered(doc, "Edge gateway validates schema, timestamps data, stores a local buffer, and forwards windows to the MARS inference API.", architecture_num)
    add_numbered(doc, "Camera/video subsystem emits detection metadata such as trespasser confidence, obstruction confidence, and camera visibility instead of streaming every frame to the model API.", architecture_num)
    add_numbered(doc, "MARS API computes track risk, weather-fusion risk, and security anomaly risk, then publishes alerts and explanations to dashboard, event bus, or maintenance workflow.", architecture_num)
    add_numbered(doc, "Operator reviews severity and recommended action; confirmed outcomes are saved for retraining and threshold tuning.", architecture_num)

    doc.add_page_break()
    add_heading(doc, "4. Practical Hardware Stack", 1)
    add_para(doc, "The following stack is feasible for a college, lab, or low-budget pilot while preserving a path to industrial-grade deployment.")
    rows = [
        ("Vibration sensing", "Budget: MPU-6050/ADXL345 for demo rig. Better pilot: ADXL355 low-noise 3-axis accelerometer. Industrial: IEPE accelerometer with signal conditioner.", "Detect vertical/lateral vibration, impacts, poor support, loosened fittings, and abnormal wheel-track interaction."),
        ("Track temperature", "Budget/pilot: waterproof DS18B20 digital probe. Industrial: RTD/PT100 with transmitter.", "Detect heat-related expansion risk and local thermal stress."),
        ("Environmental sensing", "Budget: BME280 for temperature, humidity, pressure. Pilot: tipping-bucket rain gauge plus visibility proxy. Industrial: compact weather station with rainfall, wind, visibility.", "Feed weather-track fusion: rainfall, visibility, wind, ambient temperature, fog/heat/flood flags."),
        ("Video/security", "Budget: USB or IP camera on Raspberry Pi. Pilot: PoE IP camera. Industrial: rugged thermal/network camera for low-light or fog-prone sites.", "Detect trespassers, track obstruction, platform crowding, and camera visibility issues."),
        ("Sensor node MCU", "ESP32 for low-cost Wi-Fi/Bluetooth prototypes; STM32/industrial MCU for hardened nodes.", "Samples sensors, filters noise, packages messages, and provides local watchdog/health reporting."),
        ("Edge gateway", "Raspberry Pi 5 for track/weather gateway. Jetson Orin Nano when local video AI is needed.", "Buffers data, performs lightweight preprocessing, runs camera detection, and forwards to API."),
        ("Connectivity", "Lab: Ethernet/Wi-Fi. Sparse field: LoRaWAN for low-rate sensors. Remote high-bandwidth: 4G/5G router.", "Keeps sensor data flowing while controlling cost and bandwidth."),
        ("Power", "Lab mains adapter. Field: solar panel + charge controller + LiFePO4 battery + surge protection.", "Supports remote operation and safe shutdown."),
        ("Enclosure/mounting", "IP65/IP67 enclosure, cable glands, DIN rail, vibration isolation, tamper seal.", "Protects electronics and makes maintenance repeatable."),
    ]
    add_table(doc, ["Subsystem", "Cost-effective option", "Role in MARS"], rows, [1600, 3800, 3960])

    add_heading(doc, "5. Pilot Bill of Materials Tiers", 1)
    rows = [
        ("Lab demo kit", "ESP32, MPU-6050/ADXL345, DS18B20, BME280, USB/IP camera, Raspberry Pi 5, bench power supply.", "Lowest cost; proves software flow, UI, data schema, and model behavior. Not suitable for harsh field use."),
        ("Practical field pilot", "ADXL355, DS18B20 or RTD module, rain gauge, BME280/weather module, PoE IP camera, Raspberry Pi 5 gateway, 4G router, IP65 enclosure, surge protection.", "Good balance of data quality and cost. Recommended first real pilot on controlled site."),
        ("Hardened deployment", "Industrial accelerometer, RTD/PT100, certified weather station, rugged PoE/thermal camera, Jetson edge AI, industrial gateway, UPS/solar, IP67 enclosure.", "Higher cost; use after pilot validates value and operating procedure."),
    ]
    add_table(doc, ["Tier", "Hardware set", "When to use"], rows, [1800, 4400, 3160])

    add_callout(
        doc,
        "Cost discipline",
        "Do not start with industrial hardware everywhere. Begin with inexpensive sensors on a lab rig to validate sampling rate, feature extraction, API integration, and dashboard behavior. Upgrade only the sensors whose data quality directly affects risk decisions.",
    )

    add_heading(doc, "6. Data Acquisition Plan", 1)
    rows = [
        ("Vibration", "100-1000 Hz at node; send 1-10 second windows or aggregate features", "Mean, standard deviation, peak, RMS, dominant frequency, lateral/vertical ratio"),
        ("Speed", "Per train pass or 1 Hz during simulation; source from train telemetry, GPS, axle counter, or manual feed in pilot", "Mean, max, speed ratio against permitted speed"),
        ("Track temperature", "0.1-1 Hz", "Mean, max, heat flag, deviation from regional baseline"),
        ("Weather", "1-5 minute interval", "Rainfall, visibility, ambient temperature, wind speed, flood/fog/heat flags"),
        ("Video/security", "Camera stream locally; API receives detections only", "Trespasser confidence, obstruction confidence, crowding confidence, camera visibility"),
        ("Node health", "30-60 second interval", "Battery, RSSI, packet loss, uptime, enclosure temperature"),
    ]
    add_table(doc, ["Data type", "Recommended collection rate", "Features used by MARS"], rows, [1800, 3100, 4460])

    add_heading(doc, "7. Software Integration With Current MARS Code", 1)
    add_para(doc, "The current software already has the right seams for hardware integration. Hardware should not call model internals directly; it should publish normalized events that match the existing schemas.")
    rows = [
        ("Sensor node firmware", "Reads sensor values, applies calibration, timestamps packets, sends JSON/MQTT.", "Maps to SensorEvent fields: speed, acceleration, vibration_vertical, vibration_lateral, track_temperature."),
        ("Weather station adapter", "Normalizes station output and local alerts.", "Maps to WeatherEvent fields: rainfall_mm, visibility_m, temperature_c, wind_speed_kmph, hazard_flags."),
        ("Camera edge AI", "Runs object detection locally and sends annotations or summarized confidence values.", "Maps to VideoEvent/SecurityAnomalyRequest or Hugging Face demo security inputs."),
        ("Edge gateway", "Batches events into windows, validates data, retries failed uploads, stores local buffer.", "Calls /track-risk/predict, /weather-risk/predict, /security-anomaly/predict, and /regions."),
        ("MARS backend", "Feature engineering, model inference, regional adjustment, explanation, event publishing.", "src/mars_ml_pipeline/services/app.py and feature/model modules."),
        ("Frontend/HF demo", "Operator-facing controls and visualization.", "React dashboard for local demo; Hugging Face Space for public multimodal demonstration."),
    ]
    add_table(doc, ["Layer", "Responsibility", "Current integration point"], rows, [1800, 3600, 3960])

    add_heading(doc, "8. Example Field Message Contracts", 1)
    add_para(doc, "Keep messages compact, explicit, and versioned. The gateway can translate protocol-specific payloads into these application contracts.")
    add_para(doc, "Track sensor event example:", bold=True, color=DARK_BLUE)
    add_para(doc, '{"timestamp":"2026-06-14T08:00:00Z","segment_id":"CR-SEG-001","train_id":"IR-10001","speed":96,"acceleration":0.2,"vibration_vertical":0.42,"vibration_lateral":0.30,"track_temperature":37}')
    add_para(doc, "Segment metadata example:", bold=True, color=DARK_BLUE)
    add_para(doc, '{"segment_id":"CR-SEG-001","line_id":"CR-MUMBAI","region":"CR","division":"Mumbai","asset_type":"track","age_years":22,"maintenance_score":0.68,"curvature_degree":3.2,"max_permitted_speed":105}')
    add_para(doc, "Weather event example:", bold=True, color=DARK_BLUE)
    add_para(doc, '{"region":"CR","segment_id":"CR-SEG-001","rainfall_mm":80,"visibility_m":900,"temperature_c":34,"wind_speed_kmph":70,"hazard_flags":["flood"]}')

    add_heading(doc, "9. Model Training Method", 1)
    add_para(doc, "The training workflow should combine synthetic bootstrapping, controlled pilot data, and operator-confirmed labels. Synthetic data is useful for initial demos; field labels are essential before operational use.")
    training_num = create_decimal_numbering(doc)
    add_numbered(doc, "Define the prediction target: normal, caution, high_risk. Tie labels to inspection outcomes, speed restrictions, maintenance findings, and verified incidents.", training_num)
    add_numbered(doc, "Generate baseline data using src/mars_ml_pipeline/data/generate_synthetic.py for early model behavior and CI tests.", training_num)
    add_numbered(doc, "Collect pilot data from sensor nodes and weather/video adapters. Store raw events and derived features separately.", training_num)
    add_numbered(doc, "Build track features with speed statistics, acceleration magnitude, vertical/lateral vibration statistics, temperature mean, segment age, maintenance score, curvature, and permitted speed.", training_num)
    add_numbered(doc, "Build weather-fusion features by adding rainfall, visibility, ambient temperature, wind speed, and hazard flags.", training_num)
    add_numbered(doc, "Train models with the existing scripts: train_track_risk.py and train_weather_risk.py. Export artifacts with version metadata.", training_num)
    add_numbered(doc, "Evaluate using confusion matrix, precision/recall for high-risk class, false-alarm rate, missed-event rate, and region-wise performance slices.", training_num)
    add_numbered(doc, "Deploy artifacts through MARS_TRACK_MODEL_PATH and MARS_WEATHER_MODEL_PATH. If absent, the deterministic regional rule model remains the fallback.", training_num)
    add_numbered(doc, "Continuously retrain with confirmed field outcomes, but keep a model registry and rollback path.", training_num)

    rows = [
        ("Track model", "track_sensor_events.csv", "train_track_risk.py", "mars-track-risk.joblib"),
        ("Weather fusion model", "track_sensor_events.csv + weather_events.csv", "train_weather_risk.py", "mars-weather-risk.joblib"),
        ("Security detector", "Video annotations or camera detections", "AnnotationSecurityDetector now; future YOLO/edge detector export", "security detector artifact or rules"),
    ]
    add_table(doc, ["Model", "Training data", "Training route", "Deployment artifact"], rows, [1800, 2700, 3060, 1800])

    add_heading(doc, "10. Region-Wise Practical Behavior", 1)
    rows = [
        ("CR", "Ghat sections, suburban density, monsoon disruption", "Increase vibration/weather sensitivity and reduce speed tolerance slightly."),
        ("NFR", "Hills, rivers, landslides, remote maintenance", "Highest weather and vibration sensitivity; prioritize rain, landslide, and visibility alerts."),
        ("WR/NWR", "Heat, dust, long routes, higher-speed trunk sections", "Watch thermal stress and dust visibility; speed tolerance can be higher where permitted."),
        ("ER/SR", "Waterlogging, coastal humidity, bridge approaches, cyclone exposure", "Increase weather sensitivity and corrosion/maintenance weighting."),
        ("NR", "Fog-prone plains and dense traffic", "Increase low-visibility/fog and congestion awareness."),
    ]
    add_table(doc, ["Zone group", "Typical condition", "MARS adjustment"], rows, [1500, 3600, 4260])

    add_heading(doc, "11. Deployment Phases", 1)
    rows = [
        ("Phase 0: lab proof", "1-2 weeks", "Build sensor rig, run simulator, validate API and dashboard, verify data schema.", "No field safety reliance."),
        ("Phase 1: controlled pilot", "4-8 weeks", "Install 2-5 nodes on non-critical or supervised section; compare sensor readings with manual inspection.", "Daily calibration checks; no automatic operational action."),
        ("Phase 2: extended pilot", "2-3 months", "Add weather and video/security; tune thresholds by region; collect labels for retraining.", "Human-in-the-loop advisories only."),
        ("Phase 3: hardened rollout", "After pilot acceptance", "Industrial enclosures, rugged sensors, redundancy, formal maintenance SOP.", "Integrate with operations only after safety review."),
    ]
    add_table(doc, ["Phase", "Duration", "Work", "Safety boundary"], rows, [1500, 1300, 4160, 2400])

    add_heading(doc, "12. Calibration, Reliability, and Maintenance", 1)
    add_bullet(doc, "Calibrate accelerometers on installation and after any physical maintenance.")
    add_bullet(doc, "Record mounting position, orientation, rail/fastener condition, and sensor serial number for every segment.")
    add_bullet(doc, "Use gateway buffering so outages do not destroy data continuity.")
    add_bullet(doc, "Flag stale sensors, low battery, weak signal, high packet loss, and impossible values as data-quality alerts.")
    add_bullet(doc, "Keep manual inspection feedback in the dataset; otherwise the model will learn sensor noise but not real maintenance outcomes.")
    add_bullet(doc, "Use role-based access, TLS, signed firmware, and per-device credentials before field deployment.")

    add_heading(doc, "13. Acceptance Checklist", 1)
    rows = [
        ("Hardware", "Sensor node survives vibration/heat/rain test; enclosure has strain relief and surge protection.", "Pass / Fail"),
        ("Data quality", "Timestamps are synchronized; missing data and impossible values are detected.", "Pass / Fail"),
        ("API integration", "Gateway can call /track-risk/predict, /weather-risk/predict, /security-anomaly/predict, and /regions.", "Pass / Fail"),
        ("Dashboard", "Operators can see region/division, severity, score, and recommended action.", "Pass / Fail"),
        ("Model", "High-risk recall and false-alarm rate meet pilot targets; region slices reviewed.", "Pass / Fail"),
        ("Operations", "Human review process exists before any speed restriction or field action.", "Pass / Fail"),
    ]
    add_table(doc, ["Area", "Acceptance criterion", "Status"], rows, [1700, 5960, 1700])

    add_heading(doc, "14. Recommended Minimum Viable Pilot", 1)
    add_callout(
        doc,
        "Pilot recipe",
        "Use Raspberry Pi 5 as the gateway, ESP32-based sensor nodes for low-rate telemetry, ADXL355 for vibration on the primary pilot node, DS18B20 for track temperature, BME280 plus rain gauge for weather, one PoE/IP camera for security metadata, and the existing MARS FastAPI service plus React dashboard. This is practical, cost-effective, and directly maps to the current codebase.",
    )
    add_para(doc, "Minimum working software commands:")
    add_bullet(doc, "Run backend: uvicorn mars_ml_pipeline.services.app:create_app --factory --reload --port 8000")
    add_bullet(doc, "Run regional simulator: python scripts/simulate_sensors.py --base-url http://localhost:8000")
    add_bullet(doc, "Run frontend: npm run dev from frontend/")
    add_bullet(doc, "Train track model: python -m mars_ml_pipeline.training.train_track_risk --input data/generated/track_sensor_events.csv --artifact-dir artifacts/track-risk/v0.1.0")
    add_bullet(doc, "Train weather model: python -m mars_ml_pipeline.training.train_weather_risk --sensor-input data/generated/track_sensor_events.csv --weather-input data/generated/weather_events.csv --artifact-dir artifacts/weather-risk/v0.1.0")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Appendix A: Hardware Reference Sources", 1)
    sources = [
        ("Raspberry Pi 5", "Official Raspberry Pi product page: https://www.raspberrypi.com/products/raspberry-pi-5/"),
        ("NVIDIA Jetson Orin Nano", "Official NVIDIA developer kit specifications: https://www.nvidia.com/en-us/autonomous-machines/embedded-systems/jetson-orin/nano-super-developer-kit/"),
        ("ESP32", "Espressif ESP32 product page: https://www.espressif.com/en/products/socs/esp32"),
        ("RAK LoRaWAN gateway", "RAK WisGate Edge Lite 2 product page: https://www.rakwireless.com/en-us/products/lpwan-gateways-and-concentrators/rak7268-wisgate-edge-lite-2"),
        ("ADXL355 accelerometer", "Analog Devices product/datasheet page: https://www.analog.com/en/products/adxl355.html"),
        ("BME280 environmental sensor", "Bosch Sensortec datasheet: https://www.bosch-sensortec.com/media/boschsensortec/downloads/datasheets/bst-bme280-ds002.pdf"),
        ("DS18B20 temperature probe", "Analog Devices/Maxim datasheet: https://www.analog.com/media/en/technical-documentation/data-sheets/ds18b20.pdf"),
        ("Axis thermal/network cameras", "Axis camera product categories: https://www.axis.com/products/network-cameras"),
        ("Axis Perimeter Defender", "Axis edge-based intrusion analytics: https://www.axis.com/products/axis-perimeter-defender"),
    ]
    add_table(doc, ["Reference", "Source"], sources, [2300, 7060])

    add_heading(doc, "Appendix B: Practical Notes", 1)
    add_bullet(doc, "Exact hardware brands can change during procurement. Select equivalent parts that meet the same sensing, enclosure, power, and communication requirements.")
    add_bullet(doc, "For live railway environments, installation must follow railway safety approvals, electrical isolation requirements, right-of-way rules, and maintenance authority procedures.")
    add_bullet(doc, "The MARS system should provide decision support first. Automated operational control should only be considered after formal validation, fail-safe review, and regulator/operator approval.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
